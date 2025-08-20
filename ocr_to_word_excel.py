import os
os.environ["USE_TORCH"] = "1"
try:
    from doctr.io import DocumentFile
    from doctr.models import ocr_predictor
except ImportError:
    print("Error: doctr library not found. Please install it using: pip install python-doctr")
    exit(1)

try:
    import pandas as pd
except ImportError:
    print("Error: pandas library not found. Please install it using: pip install pandas")
    exit(1)

try:
    from docx import Document
except ImportError:
    print("Error: python-docx library not found. Please install it using: pip install python-docx")
    exit(1)

from datetime import datetime
import re

PDF_PATH = "invoices/your_invoice.pdf"
OUTPUT_FOLDER = "extracted_data"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

model = ocr_predictor(pretrained=True)
doc = DocumentFile.from_pdf(PDF_PATH)
result = model(doc)

# Keywords for company name
company_keywords = ["GARMENTS", "TRADING", "LLC", "COLLECTIONS", "TEXTILES", "COMPANY", "CORPORATION", "EST", "SUPPLIERS", "INDUSTRIAL", "UNIFORMS"]

def extract_company_name(lines):
    company_keywords = [
        "LLC", "L.L.C", "TRADING", "GARMENTS", "COMPANY", "COLLECTIONS", "TEXTILES",
        "CORPORATION", "EST", "SUPPLIERS", "INDUSTRIAL", "UNIFORMS", "AREA"
    ]
    # 1. Look for a line with a company keyword in the first 30 lines
    for line in lines[:30]:
        if any(kw in line.upper() for kw in company_keywords):
            return line.strip()
    # 2. Fallback: first long line not containing common headers
    for line in lines[:30]:
        if len(line) > 6 and not any(x in line.upper() for x in [
            "INVOICE", "DATE", "TOTAL", "AMOUNT", "VAT", "TRN", "BILL", "NUMBER", "QUANTITY", "ADDRESS"
        ]):
            return line.strip()
    return "Not Found"

def extract_invoice_number(lines, text):
    # 1. If a line contains 'Invoice No.', check that line and next 2 lines for invoice pattern
    for i, line in enumerate(lines[:15]):
        if 'INVOICE NO' in line.upper():
            for j in range(i, min(i+3, len(lines))):
                m = re.search(r'[A-Z]{2,4}/\d{4,}/?\d{0,6}|[A-Z]{2,4}/\d{6,}', lines[j])
                if m:
                    return m.group(0)
    # 2. Search first 15 lines for invoice pattern
    for line in lines[:15]:
        m = re.search(r'[A-Z]{2,4}/\d{4,}/?\d{0,6}|[A-Z]{2,4}/\d{6,}', line)
        if m:
            return m.group(0)
    return "Not Found"

def extract_date(lines, text):
    for line in lines:
        if 'DATED' in line.upper() or 'DATE' in line.upper():
            m = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}-[A-Za-z]{3}-\d{4})', line)
            if m:
                return m.group(1)
    m = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}-[A-Za-z]{3}-\d{4})', text)
    if m:
        return m.group(1)
    return "Not Found"

def extract_trns(text):
    trns = re.findall(r'TRN\s*:?\s*(\d{9,15})', text, re.IGNORECASE)
    seller = trns[0] if len(trns) > 0 else "Not Found"
    buyer = trns[1] if len(trns) > 1 else seller
    return seller, buyer

def parse_number(val):
    try:
        return float(val.replace(',', '').replace('AED', '').strip())
    except Exception:
        return None

def extract_amount_before_tax(lines):
    """
    Extract amount before tax (subtotal) by looking for total amounts before VAT
    """
    for i, line in enumerate(lines):
        line_upper = line.upper()
        if 'TOTAL' in line_upper and 'VAT' not in line_upper and 'IN WORDS' not in line_upper:
            nums = re.findall(r'[\d,.]+', line)
            if nums:
                return nums[-1]
    
    # Fallback: look for numbers before VAT section
    for i in range(len(lines)-1, -1, -1):
        line_upper = lines[i].upper()
        if 'VAT' in line_upper:
            # Look for numbers in previous lines
            for j in range(i-1, max(0, i-5), -1):
                prev_line = lines[j].strip()
                nums = re.findall(r'[\d,.]+', prev_line)
                if nums:
                    return nums[-1]
    
    return "Not Found"

def extract_vat_amount(lines):
    """
    Extract VAT amount by looking for values exactly in front of or below "VAT 5%" patterns, prioritizing 'AED' values and the smallest number after 'VAT'.
    """
    def safe_float(num_str):
        try:
            return float(num_str.replace(',', ''))
        except (ValueError, AttributeError):
            return float('inf')

    for i, line in enumerate(lines):
        line_upper = line.upper()
        # Look for VAT 5% patterns
        if 'VAT' in line_upper and ('5%' in line_upper or '5' in line_upper):
            # 1. Prefer 'AED' value on the same line
            m = re.search(r'AED\s*([\d,.]+)', line_upper)
            if m:
                return m.group(1)
            # 2. Prefer 'AED' value on the next line
            if i < len(lines) - 1:
                m2 = re.search(r'AED\s*([\d,.]+)', lines[i+1].upper())
                if m2:
                    return m2.group(1)
            # 3. Otherwise, pick the smallest number on the same or next line (not 5)
            nums = re.findall(r'[\d,.]+', line)
            if i < len(lines) - 1:
                nums += re.findall(r'[\d,.]+', lines[i+1])
            valid_nums = [num for num in nums if num.replace(',', '') not in ['5', '5.0', '5.00'] and safe_float(num) != float('inf')]
            if valid_nums:
                return min(valid_nums, key=safe_float)
    # Fallback: previous logic
    for i, line in enumerate(lines):
        line_upper = line.upper()
        if 'VAT' in line_upper:
            nums = re.findall(r'[\d,.]+', line)
            valid_nums = [num for num in nums if num.replace(',', '') not in ['5', '5.0', '5.00'] and safe_float(num) != float('inf')]
            if valid_nums:
                return min(valid_nums, key=safe_float)
            # Check 2 lines above and below for VAT amount
            for offset in [-2, -1, 1, 2]:
                check_idx = i + offset
                if 0 <= check_idx < len(lines):
                    check_line = lines[check_idx].strip()
                    check_nums = re.findall(r'[\d,.]+', check_line)
                    valid_nums = [num for num in check_nums if num.replace(',', '') not in ['5', '5.0', '5.00'] and safe_float(num) != float('inf')]
                    if valid_nums:
                        return min(valid_nums, key=safe_float)
    return "Not Found"

def extract_subtotal(lines):
    summary_keywords = ['VAT', 'GRAND TOTAL', 'TOTAL AMOUNT', 'NET AMOUNT', 'AMOUNT PAYABLE', 'TOTAL AED', 'TOTAL']
    summary_idx = None
    for i in range(len(lines)-1, -1, -1):
        l = lines[i].upper()
        if any(kw in l for kw in summary_keywords):
            summary_idx = i
            break
    # 2. Look for 'Subtotal' label before summary
    if summary_idx:
        for i in range(summary_idx-1, -1, -1):
            if 'SUBTOTAL' in lines[i].upper():
                nums = re.findall(r'[\d,.]+', lines[i])
                if nums:
                    return nums[-1]
    # 3. Fallback: last number before summary section
    stop_idx = summary_idx if summary_idx is not None else len(lines)
    for i in range(stop_idx-1, -1, -1):
        l = lines[i].upper()
        if not any(kw in l for kw in summary_keywords):
            nums = re.findall(r'[\d,.]+', lines[i])
            if nums:
                return nums[-1]
    return "Not Found"

def extract_total_amount(lines):
    for i, line in enumerate(lines):
        if ('GRAND TOTAL' in line.upper() or 'TOTAL AED' in line.upper() or 'TOTAL AMOUNT' in line.upper() or 'NET AMOUNT' in line.upper() or 'AMOUNT PAYABLE' in line.upper()):
            m = re.search(r'AED\s*([\d,.]+)', line)
            if m:
                return m.group(1)
            nums = re.findall(r'[\d,.]+', line)
            if nums:
                return nums[-1]
        # Fallback: 'Total' with AED
        if 'TOTAL' in line.upper() and 'AED' in line.upper():
            m = re.search(r'AED\s*([\d,.]+)', line)
            if m:
                return m.group(1)
    # Fallback: last number in the last 10 lines
    for line in lines[::-1][:10]:
        nums = re.findall(r'[\d,.]+', line)
        if nums:
            return nums[-1]
    return "Not Found"

def extract_amounts(lines):
    subtotal = total = "Not Found"
    
    exclude_keywords = [
        'VAT', 'IN WORDS', 'DISCOUNT', 'DIRHAM', 'PAID', 'ROUND', 'FILS', 'LESS', 'CHARGEABLE',
        'VAT AMOUNT', 'TAX AMOUNT', 'AMOUNT IN WORDS', 'DECLARATION', 'SIGNATURE', 'AUTHORIZED',
        'AUTHORISED', 'THANK YOU', 'SUBTOTAL', 'BALANCE', 'DUE', 'RECEIVED', 'CASH', 'BANK',
        'ACCOUNT', 'IBAN', 'NO.', 'A/C', 'ALC', 'BRANCH', 'SWIFT', 'CODE', 'CHEQUE', 'PAYEE',
        'TRN', 'TRADING', 'SUPPLIER', 'BUYER', 'INVOICE', 'DATE', 'ORDER', 'REFERENCE',
        'DELIVERY', 'NOTE', 'QUANTITY', 'RATE', 'DESCRIPTION', 'ITEM', 'PCS', 'QTY', 'UNIT', 'PRICE',
        'AMOUNT', 'TOTAL VAT', 'VAT 5%', 'VAT 5', 'VAT5', 'VAT5%'
    ]
    
    candidate_amounts = []
    for i, line in enumerate(lines):
        line_upper = line.upper()
        # Only consider lines with 'TOTAL' and not with any exclude keywords
        if 'TOTAL' in line_upper and not any(keyword in line_upper for keyword in exclude_keywords):
            # Find all numbers in this line
            nums = re.findall(r'[\d,.]+', line)
            for num in nums:
                try:
                    value = float(num.replace(',', ''))
                    candidate_amounts.append(value)
                except Exception:
                    pass
            # Also check the next line for numbers (sometimes total is on the next line)
            if i+1 < len(lines):
                next_line = lines[i+1]
                next_line_upper = next_line.upper()
                if not any(keyword in next_line_upper for keyword in exclude_keywords):
                    nums = re.findall(r'[\d,.]+', next_line)
                    for num in nums:
                        try:
                            value = float(num.replace(',', ''))
                            candidate_amounts.append(value)
                        except Exception:
                            pass
    if candidate_amounts:
        total = f"AED {max(candidate_amounts):,.2f}"
    return subtotal, total

def extract_invoice_table(ocr_lines, invoice_number=None, page_idx=None):
    header_idx = None
    for i, line in enumerate(ocr_lines):
        if (re.search(r'qty|quantity', line, re.I) and
            re.search(r'rate', line, re.I) and
            re.search(r'amount', line, re.I)):
            header_idx = i
            break
    if header_idx is None:
        print("No table header found.")
        return pd.DataFrame()
    items = []
    for line in ocr_lines[header_idx+1:]:
        if re.search(r'total|vat|grand total|net amount|amount payable', line, re.I):
            break
        cols = re.split(r'\s{2,}|\t', line)
        if len(cols) >= 4:
            try:
                qty = float(cols[0].replace(',', '')) if cols[0].replace(',', '').replace('.', '').isdigit() else float('nan')
                desc = cols[1]
                rate = float(cols[2].replace(',', '')) if cols[2].replace(',', '').replace('.', '').isdigit() else float('nan')
                amount = float(cols[-1].replace(',', '')) if cols[-1].replace(',', '').replace('.', '').isdigit() else float('nan')
                valid = (not pd.isna(qty) and not pd.isna(rate) and not pd.isna(amount) and abs(qty * rate - amount) < 0.05)
                items.append({
                    'Invoice Number': invoice_number if invoice_number else '',
                    'Page': page_idx + 1 if page_idx is not None else '',
                    'Quantity': qty,
                    'Description': desc,
                    'Rate': rate,
                    'Amount': amount,
                    'Valid': valid
                })
            except Exception:
                continue
    df = pd.DataFrame(items)
    if not df.empty:
        df['Quantity'] = pd.to_numeric(df['Quantity'], errors='coerce')
        df['Rate'] = pd.to_numeric(df['Rate'], errors='coerce')
        df['Amount'] = pd.to_numeric(df['Amount'], errors='coerce')
        total_qty = df['Quantity'].sum()
        total_amt = df['Amount'].sum()
        print(f"Total Quantity: {total_qty}, Total Amount: {total_amt}")
        print(df)
    else:
        print("No valid items found.")
    return df

def safe_int(val):
    try:
        if val in (None, 'Not Found'):
            return 1
        if isinstance(val, int):
            return val
        if isinstance(val, float):
            return int(val)
        if isinstance(val, str) and val.isdigit():
            return int(val)
        return 1
    except Exception:
        return 1

def extract_items_from_pdf(pdf_path):
    model = ocr_predictor(pretrained=True)
    doc = DocumentFile.from_pdf(pdf_path)
    result = model(doc)
    rows = []
    all_items = []
    for page_idx, page in enumerate(result.pages):
        page_lines = []
        for block in page.blocks:
            for line in block.lines:
                line_text = " ".join([word.value for word in line.words])
                page_lines.append(line_text)
        page_text = " ".join(page_lines)
        row = {}
        row['Page'] = int(page_idx + 1)
        row['Company Name'] = extract_company_name(page_lines)
        row['Invoice Number'] = extract_invoice_number(page_lines, page_text)
        row['Date'] = extract_date(page_lines, page_text)
        seller_trn, buyer_trn = extract_trns(page_text)
        row['Seller TRN'] = seller_trn
        row['Buyer TRN'] = buyer_trn
        _, total = extract_amounts(page_lines)
        vat_amount = extract_vat_amount(page_lines)
        print(f"[DEBUG] Page {page_idx+1}: VAT Amount: {vat_amount}, Total Amount: {total}")
        row['VAT Amount'] = vat_amount
        row['Total Amount'] = total
        rows.append(row)
        item_df = extract_invoice_table(page_lines, invoice_number=row['Invoice Number'], page_idx=page_idx)
        if not item_df.empty:
            all_items.append(item_df)
    if all_items:
        items_df = pd.concat(all_items, ignore_index=True)
    else:
        items_df = pd.DataFrame()
    summary_row = rows[0] if rows else {}
    # Build a cleaned summary row with correct types
    cleaned_summary_row = {}
    for k, v in summary_row.items():
        if k == 'Page':
            cleaned_summary_row[k] = safe_int(v)
        elif v in (None, 'Not Found'):
            cleaned_summary_row[k] = None
        elif isinstance(v, float) and pd.isna(v):
            cleaned_summary_row[k] = None
        elif isinstance(v, (float, int)):
            cleaned_summary_row[k] = float(v)
        else:
            cleaned_summary_row[k] = str(v)
    return cleaned_summary_row, items_df

if __name__ == '__main__':
    PDF_PATH = "invoices/your_invoice.pdf"
    OUTPUT_FOLDER = "extracted_data"
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    model = ocr_predictor(pretrained=True)
    doc = DocumentFile.from_pdf(PDF_PATH)
    result = model(doc)
    rows = []
    all_items = []
    for page_idx, page in enumerate(result.pages):
        page_lines = []
        for block in page.blocks:
            for line in block.lines:
                line_text = " ".join([word.value for word in line.words])
                page_lines.append(line_text)
        page_text = " ".join(page_lines)
        print(f"\n--- OCR LINES (Last 15) for Page {page_idx+1} ---")
        for i, l in enumerate(page_lines[-15:]):
            print(f"{len(page_lines)-15+i+1}: {l}")
        print("--- END OCR LINES ---\n")
        row = {"Page": page_idx + 1}
        row['Company Name'] = extract_company_name(page_lines)
        row['Invoice Number'] = extract_invoice_number(page_lines, page_text)
        row['Date'] = extract_date(page_lines, page_text)
        seller_trn, buyer_trn = extract_trns(page_text)
        row['Seller TRN'] = seller_trn
        row['Buyer TRN'] = buyer_trn
        _, total = extract_amounts(page_lines)
        vat_amount = extract_vat_amount(page_lines)
        print(f"[DEBUG] Page {page_idx+1}: VAT Amount: {vat_amount}, Total Amount: {total}")
        row['VAT Amount'] = vat_amount
        row['Total Amount'] = total
        rows.append(row)
        item_df = extract_invoice_table(page_lines, invoice_number=row['Invoice Number'], page_idx=page_idx)
        if not item_df.empty:
            item_excel_path = os.path.join(OUTPUT_FOLDER, f"items_table_page_{page_idx+1}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            item_df.to_excel(item_excel_path, index=False)
            print(f"Saved item table to Excel: {item_excel_path}")
            all_items.append(item_df)
    excel_path = os.path.join(OUTPUT_FOLDER, f"multi_invoice_table_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    df = pd.DataFrame(rows)
    df.to_excel(excel_path, index=False)
    print(f"Saved multi-invoice table to Excel: {excel_path}")
    docx_path = os.path.join(OUTPUT_FOLDER, f"multi_invoice_table_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    docx = Document()
    docx.add_heading('Multi-Invoice Extraction Table', 0)
    docx.add_paragraph(f"File: {PDF_PATH}")
    docx.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    docx.add_paragraph("")
    table = docx.add_table(rows=1, cols=len(rows[0]))
    hdr_cells = table.rows[0].cells
    for i, key in enumerate(rows[0].keys()):
        hdr_cells[i].text = key
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = str(value)
    docx.save(docx_path)
    print(f"Saved multi-invoice table to Word: {docx_path}")
    if all_items:
        master_df = pd.concat(all_items, ignore_index=True)
        master_path = os.path.join(OUTPUT_FOLDER, f"all_items_master_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        master_df.to_excel(master_path, index=False)
        print(f"Saved master item table to Excel: {master_path}") 