import re
import os
import cv2
import numpy as np
import pandas as pd
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from doctr.io import DocumentFile
from doctr.models import ocr_predictor
from docx import Document
from docx.shared import Inches
import warnings
warnings.filterwarnings('ignore')

# Try to import PyPDF2 for text-based PDFs
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

os.environ["USE_TORCH"] = "1"

class InvoiceDataExtractor:
    """
    Comprehensive invoice data extraction using Doctr OCR
    Extracts: Company Name, Invoice Number, Date, TRN, Quantity, Amount, VAT
    """
    
    def __init__(self):
        """Initialize the OCR model and patterns"""
        self.model = ocr_predictor(pretrained=True)
        
        # Regex patterns for different data types
        self.patterns = {
            'invoice_number': [
                r'invoice\s*#?\s*:?\s*([A-Z0-9\-_]+)',
                r'invoice\s*number\s*:?\s*([A-Z0-9\-_]+)',
                r'inv\s*:?\s*([A-Z0-9\-_]+)',
                r'#\s*([A-Z0-9\-_]+)',
                r'bill\s*#\s*:?\s*([A-Z0-9\-_]+)',
                r'order\s*([A-Z0-9\-_]+)'
            ],
            'date': [
                r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
                r'(\d{4}[/\-]\d{1,2}[/\-]\d{1,2})',
                r'(\d{1,2}\s+\w+\s+\d{4})',
                r'date\s*:?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',
                r'order\s*date\s*:?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})'
            ],
            'trn': [
                r'trn\s*:?\s*(\d{9,15})',
                r'tax\s*registration\s*number\s*:?\s*(\d{9,15})',
                r'vat\s*number\s*:?\s*(\d{9,15})',
                r'(\d{9,15})'
            ],
            'amount': [
                r'total\s*:?\s*([\d,]+\.?\d*)\s*([A-Z]{3})',
                r'amount\s*:?\s*([\d,]+\.?\d*)\s*([A-Z]{3})',
                r'([\d,]+\.?\d*)\s*([A-Z]{3})',
                r'([\d,]+\.?\d*)',
                r'grand\s*total\s*:?\s*([\d,]+\.?\d*)\s*([A-Z]{3})',
                r'final\s*amount\s*:?\s*([\d,]+\.?\d*)\s*([A-Z]{3})'
            ],
            'vat_amount': [
                r'vat\s*:?\s*([\d,]+\.?\d*)',
                r'tax\s*:?\s*([\d,]+\.?\d*)',
                r'gst\s*:?\s*([\d,]+\.?\d*)'
            ],
            'quantity': [
                r'quantity\s*:?\s*([\d,]+\.?\d*)\s*m',
                r'qty\s*:?\s*([\d,]+\.?\d*)\s*m',
                r'([\d,]+\.?\d*)\s*meters',
                r'([\d,]+\.?\d*)\s*m'
            ]
        }
    
    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict]:
        """
        Extract text from PDF using both PyPDF2 (for text-based) and Doctr OCR (for image-based)
        """
        try:
            # First try PyPDF2 for text-based PDFs
            if HAS_PYPDF2:
                text_data = self._extract_text_with_pypdf2(pdf_path)
                if text_data:
                    return text_data
            
            # Fall back to Doctr OCR for image-based PDFs
            return self._extract_text_with_doctr(pdf_path)
            
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return []
    
    def _extract_text_with_pypdf2(self, pdf_path: str) -> List[Dict]:
        """Extract text from text-based PDF using PyPDF2"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_data = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        # Split text into words and create structured data
                        words = text.split()
                        for i, word in enumerate(words):
                            text_data.append({
                                'text': word,
                                'confidence': 1.0,  # High confidence for text-based PDFs
                                'bbox': None
                            })
                
                return text_data
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
            return []
    
    def _extract_text_with_doctr(self, pdf_path: str) -> List[Dict]:
        """Extract text from image-based PDF using Doctr OCR"""
        try:
            # Load document
            doc = DocumentFile.from_pdf(pdf_path)
            
            # Perform OCR
            result = self.model(doc)
            
            # Extract text with positions
            extracted_data = []
            for page in result.pages:
                for block in page.blocks:
                    for line in block.lines:
                        for word in line.words:
                            extracted_data.append({
                                'text': word.value,
                                'confidence': word.confidence,
                                'bbox': word.geometry
                            })
            
            return extracted_data
        except Exception as e:
            print(f"Doctr OCR extraction failed: {e}")
            return []
    
    def find_company_name(self, text_data: List[Dict]) -> str:
        """
        Extract company name using uppercase heuristic
        """
        company_candidates = []
        
        for item in text_data:
            text = item['text'].strip()
            confidence = item['confidence']
            
            # Look for uppercase text that could be company names
            if (len(text) > 3 and 
                text.isupper() and 
                confidence > 0.7 and
                not any(char.isdigit() for char in text) and
                not text.startswith('INVOICE') and
                not text.startswith('DATE') and
                not text.startswith('TOTAL') and
                not text.startswith('BILL')):
                
                company_candidates.append((text, confidence))
        
        # Return the highest confidence company name
        if company_candidates:
            return max(company_candidates, key=lambda x: x[1])[0]
        
        return "Not Found"
    
    def extract_field(self, text_data: List[Dict], field_patterns: List[str]) -> str:
        """
        Extract specific field using regex patterns
        """
        full_text = ' '.join([item['text'] for item in text_data])
        
        for pattern in field_patterns:
            matches = re.findall(pattern, full_text, re.IGNORECASE)
            if matches:
                return matches[0] if isinstance(matches[0], str) else ' '.join(matches[0])
        
        return "Not Found"
    
    def extract_invoice_data(self, pdf_path: str) -> Dict:
        """
        Extract all required data from invoice PDF
        """
        print(f"Processing: {pdf_path}")
        
        # Extract text from PDF
        text_data = self.extract_text_from_pdf(pdf_path)
        
        if not text_data:
            return {"error": "No text extracted from PDF"}
        
        # Extract all fields
        extracted_data = {
            'file_name': os.path.basename(pdf_path),
            'company_name': self.find_company_name(text_data),
            'invoice_number': self.extract_field(text_data, self.patterns['invoice_number']),
            'date': self.extract_field(text_data, self.patterns['date']),
            'seller_trn': self.extract_field(text_data, self.patterns['trn']),
            'buyer_trn': self.extract_field(text_data, self.patterns['trn']),  # May need refinement
            'total_quantity_meters': self.extract_field(text_data, self.patterns['quantity']),
            'total_amount': self.extract_field(text_data, self.patterns['amount']),
            'vat_amount': self.extract_field(text_data, self.patterns['vat_amount']),
            'extraction_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        return extracted_data
    
    def save_to_excel(self, data_list: List[Dict], output_path: str):
        """
        Save extracted data to Excel file
        """
        try:
            df = pd.DataFrame(data_list)
            
            # Create Excel writer
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Invoice Data', index=False)
                
                # Auto-adjust column widths
                worksheet = writer.sheets['Invoice Data']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            print(f"Data saved to Excel: {output_path}")
            
        except Exception as e:
            print(f"Error saving to Excel: {e}")
    
    def save_to_word(self, data_list: List[Dict], output_path: str):
        """
        Save extracted data to Word document
        """
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Invoice Data Extraction Report', 0)
            title.alignment = 1  # Center alignment
            
            # Add summary
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Total invoices processed: {len(data_list)}")
            doc.add_paragraph("")
            
            # Add table
            table = doc.add_table(rows=1, cols=len(data_list[0].keys()))
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, key in enumerate(data_list[0].keys()):
                header_cells[i].text = key.replace('_', ' ').title()
            
            # Add data rows
            for data in data_list:
                row_cells = table.add_row().cells
                for i, value in enumerate(data.values()):
                    row_cells[i].text = str(value)
            
            # Save document
            doc.save(output_path)
            print(f"Data saved to Word: {output_path}")
            
        except Exception as e:
            print(f"Error saving to Word: {e}")
    
    def process_invoices(self, input_folder: str, output_folder: str = "output"):
        """
        Process all PDF invoices in a folder
        """
        # Create output folder
        os.makedirs(output_folder, exist_ok=True)
        
        # Find all PDF files
        pdf_files = [f for f in os.listdir(input_folder) if f.lower().endswith('.pdf')]
        
        if not pdf_files:
            print("No PDF files found in the input folder")
            return
        
        # Process each PDF
        all_data = []
        for pdf_file in pdf_files:
            pdf_path = os.path.join(input_folder, pdf_file)
            data = self.extract_invoice_data(pdf_path)
            
            if 'error' not in data:
                all_data.append(data)
                print(f"✓ Processed: {pdf_file}")
            else:
                print(f"✗ Failed: {pdf_file} - {data['error']}")
        
        if all_data:
            # Save to Excel
            excel_path = os.path.join(output_folder, f"invoice_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            self.save_to_excel(all_data, excel_path)
            
            # Save to Word
            word_path = os.path.join(output_folder, f"invoice_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            self.save_to_word(all_data, word_path)
            
            print(f"\nProcessing complete! {len(all_data)} invoices processed.")
            print(f"Files saved to: {output_folder}")
        else:
            print("No data extracted from any invoices")

def main():
    """
    Main function to demonstrate usage
    """
    # Initialize extractor
    extractor = InvoiceDataExtractor()
    
    # Set input and output folders
    input_folder = "invoices"  # Put your PDF invoices here
    output_folder = "extracted_data"
    
    # Process invoices
    extractor.process_invoices(input_folder, output_folder)

if __name__ == "__main__":
    main() 