import os
os.environ["USE_TORCH"] = "1"
try:
    from doctr.io import DocumentFile
    from doctr.models import ocr_predictor
except ImportError:
    print("Error: doctr library not found. Please install it using: pip install python-doctr")
    exit(1)

try:
    import pandas as pd
except ImportError:
    print("Error: pandas library not found. Please install it using: pip install pandas")
    exit(1)

try:
    from docx import Document
except ImportError:
    print("Error: python-docx library not found. Please install it using: pip install python-docx")
    exit(1)

from datetime import datetime
import re
from pathlib import Path

PDF_PATH = "invoices/your_invoice.pdf"
OUTPUT_FOLDER = "extracted_data"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Global variable for model (will be loaded lazily)
_model = None

def get_model():
    """Lazy load the OCR model only when needed"""
    global _model
    if _model is None:
        print("Loading OCR model...")
        _model = ocr_predictor(pretrained=True)
        print("OCR model loaded successfully!")
    return _model

def process_invoice(pdf_path=PDF_PATH):
    """Main function to process the invoice. Supports PDF and image files (png/jpg/jpeg)."""
    try:
        print(f"Processing invoice: {pdf_path}")
        
        # Check if file exists
        if not os.path.exists(pdf_path):
            print(f"Error: PDF file not found at {pdf_path}")
            return False
            
        # Load model and process
        model = get_model()
        file_ext = Path(pdf_path).suffix.lower()
        if file_ext in [".png", ".jpg", ".jpeg"]:
            doc = DocumentFile.from_images([pdf_path])
        else:
            doc = DocumentFile.from_pdf(pdf_path)
        print("Running OCR...")
        result = model(doc)
        print("OCR completed!")
        
        # Structured export for layout-aware parsing
        export_data = result.export()
        # Build per-page texts for keyword-only fallbacks
        page_texts = extract_page_texts(result)
        if not page_texts:
            # Fallback: single combined text
            combined_text = result.render()
            page_texts = [combined_text]
        
        rows = []
        for idx, page_text in enumerate(page_texts, start=1):
            lines = page_text.split('\n')
            company_name = extract_company_name(lines)
            invoice_number = extract_invoice_number(lines, page_text)
            date = extract_date(lines, page_text)
            seller_trn, buyer_trn = extract_trns(page_text)
            # Prefer layout-aware extraction using word coordinates
            page_dict = export_data.get('pages', [])[idx - 1] if export_data.get('pages') else None
            subtotal_layout, vat_layout, total_layout = extract_amounts_layout(page_dict)

            # Fallbacks to text-only heuristics
            vat_amount = vat_layout if vat_layout != "Not Found" else extract_vat_amount(lines)
            total_amount = total_layout if total_layout != "Not Found" else extract_total_amount(lines)
            rows.append({
                'Page': idx,
                'Company Name': company_name,
                'Invoice Number': invoice_number,
                'Date': date,
                'Seller TRN': seller_trn,
                'Buyer TRN': buyer_trn,
                'VAT Amount': vat_amount,
                'Total Amount': total_amount,
            })
        
        print("Data extracted successfully!")
        
        # Save tabular outputs
        save_table_to_excel(rows)
        save_table_to_word(rows)
        
        return True
        
    except Exception as e:
        print(f"Error processing invoice: {e}")
        return False

def extract_page_texts(result) -> list:
    """Extract plain text per page from doctr result.export() structure."""
    try:
        data = result.export()
    except Exception:
        return []
    page_texts = []
    for page in data.get('pages', []):
        lines_text = []
        for block in page.get('blocks', []):
            for line in block.get('lines', []):
                words = [w.get('value', '') for w in line.get('words', [])]
                if words:
                    lines_text.append(' '.join(words))
        page_texts.append('\n'.join(lines_text))
    return page_texts

# Keywords for company name
company_keywords = ["GARMENTS", "TRADING", "LLC", "COLLECTIONS", "TEXTILES", "COMPANY", "CORPORATION", "EST", "SUPPLIERS", "INDUSTRIAL", "UNIFORMS"]

def extract_company_name(lines):
    company_keywords = [
        "LLC", "L.L.C", "TRADING", "GARMENTS", "COMPANY", "COLLECTIONS", "TEXTILES",
        "CORPORATION", "EST", "SUPPLIERS", "INDUSTRIAL", "UNIFORMS", "AREA"
    ]
    # 1. Look for a line with a company keyword in the first 30 lines
    for line in lines[:30]:
        if any(kw in line.upper() for kw in company_keywords):
            return line.strip()
    # 2. Fallback: first long line not containing common headers
    for line in lines[:30]:
        if len(line) > 6 and not any(x in line.upper() for x in [
            "INVOICE", "DATE", "TOTAL", "AMOUNT", "VAT", "TRN", "BILL", "NUMBER", "QUANTITY", "ADDRESS"
        ]):
            return line.strip()
    return "Not Found"

def extract_invoice_number(lines, text):
    # 1. If a line contains 'Invoice No.', check that line and next 2 lines for invoice pattern
    for i, line in enumerate(lines[:15]):
        if 'INVOICE NO' in line.upper():
            for j in range(i, min(i+3, len(lines))):
                m = re.search(r'[A-Z]{2,4}/\d{4,}/?\d{0,6}|[A-Z]{2,4}/\d{6,}', lines[j])
                if m:
                    return m.group(0)
    # 2. Search first 15 lines for invoice pattern
    for line in lines[:15]:
        m = re.search(r'[A-Z]{2,4}/\d{4,}/?\d{0,6}|[A-Z]{2,4}/\d{6,}', line)
        if m:
            return m.group(0)
    return "Not Found"

def extract_date(lines, text):
    for line in lines:
        if 'DATED' in line.upper() or 'DATE' in line.upper():
            m = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}-[A-Za-z]{3}-\d{4})', line)
            if m:
                return m.group(1)
    m = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}-[A-Za-z]{3}-\d{4})', text)
    if m:
        return m.group(1)
    return "Not Found"

def extract_trns(text):
    trns = re.findall(r'TRN\s*:?\s*(\d{9,15})', text, re.IGNORECASE)
    seller = trns[0] if len(trns) > 0 else "Not Found"
    buyer = trns[1] if len(trns) > 1 else seller
    return seller, buyer

def parse_number(val):
    try:
        s = val.upper().replace('AED', '')
        s = re.sub(r'[^0-9.,-]', '', s)
        # Prefer dot as decimal separator
        if s.count(',') > 1 and '.' not in s:
            s = s.replace(',', '')
        else:
            s = s.replace(',', '')
        return float(s)
    except Exception:
        return None

def extract_amounts_layout(page: dict) -> tuple[str, str, str]:
    """Layout-aware extraction of (subtotal, vat, total) by reading horizontally.

    Strategy:
    - Build a list of words with text and bbox (x0,y0,x1,y1) normalized 0..1.
    - Search for anchors: ['SUBTOTAL','SUB-TOTAL','AMOUNT BEFORE TAX','BEFORE VAT'], ['VAT'],
      and final total anchors ['GRAND TOTAL','TOTAL AMOUNT','AMOUNT PAYABLE','NET PAYABLE','INVOICE TOTAL','TOTAL'].
    - For each anchor, look to the right on approximately the same line (y overlap) and pick the rightmost numeric token
      that looks like currency/amount.
    - Return strings (formatted to 2 decimals) or "Not Found".
    """
    if not page:
        return "Not Found", "Not Found", "Not Found"

    def iter_words():
        for block in page.get('blocks', []):
            for line in block.get('lines', []):
                for word in line.get('words', []):
                    val = word.get('value', '')
                    box = word.get('geometry', [[0,0],[0,0]])
                    # geometry: [[x0,y0],[x1,y1]]
                    try:
                        (x0,y0),(x1,y1) = box
                    except Exception:
                        x0=y0=x1=y1=0.0
                    yield {
                        'text': val,
                        'x0': float(x0), 'y0': float(y0), 'x1': float(x1), 'y1': float(y1),
                        'yc': float(y0+y1)/2.0
                    }

    words = list(iter_words())
    if not words:
        return "Not Found", "Not Found", "Not Found"

    # Uppercase helper available to all nested functions
    text_upper = lambda s: (s or '').upper()

    def is_amount_token(t: str) -> bool:
        return bool(re.search(r'^(?:AED\s*)?[\d,.]+(?:\s*AED)?$', t, re.IGNORECASE))

    def parse_amount_token(t: str):
        m = re.search(r'[\d,.]+', t)
        if not m:
            return None
        return parse_number(m.group(0))

    def get_amount_column_bounds() -> tuple[float, float] | None:
        # 1) Try header 'AMOUNT'
        header = None
        for w in words:
            if 'AMOUNT' in text_upper(w['text']) and len(w['text']) <= 10:
                header = w
                break
        if header:
            # assume amounts are to the right of the header start
            return max(0.0, header['x0'] - 0.02), 1.0
        # 2) Infer from numeric tokens clustered on the right
        numeric_tokens = [w for w in words if re.match(r'^(?:AED\s*)?[\d,.]+(?:\s*AED)?$', w['text'], re.IGNORECASE)]
        if not numeric_tokens:
            return None
        max_x0 = max(w['x0'] for w in numeric_tokens)
        # set a band to capture the rightmost column
        return max(0.0, max_x0 - 0.2), 1.0

    amount_bounds = get_amount_column_bounds()

    def in_amount_column(w) -> bool:
        if amount_bounds is None:
            return True
        left, right = amount_bounds
        return left <= w['x0'] <= right

    def nearest_right_amount(anchor, max_dx=0.6, y_tol=0.02):
        ax1 = anchor['x1']
        ay = anchor['yc']
        candidates = []
        for w in words:
            # right side and similar y, constrained to amount column if known
            if w['x0'] >= ax1 and abs(w['yc'] - ay) <= y_tol and (w['x0'] - ax1) <= max_dx and in_amount_column(w):
                if is_amount_token(w['text']):
                    val = parse_amount_token(w['text'])
                    if val is not None:
                        candidates.append((w['x0'], val))
        if not candidates:
            return None
        # choose the rightmost (largest x0)
        candidates.sort(key=lambda t: t[0], reverse=True)
        return candidates[0][1]

    def find_anchor(predicate):
        for w in words:
            if predicate(w['text']):
                return w
        return None

    # Anchors
    subtotal_anchor = find_anchor(lambda t: any(k in text_upper(t) for k in ['SUBTOTAL','SUB-TOTAL','AMOUNT BEFORE TAX','BEFORE VAT']))
    total_anchor = find_anchor(lambda t: any(k in text_upper(t) for k in ['GRAND TOTAL','TOTAL AMOUNT','AMOUNT PAYABLE','NET PAYABLE','INVOICE TOTAL']))
    if total_anchor is None:
        # fallback to a plain TOTAL that is not quantity related
        total_anchor = find_anchor(lambda t: 'TOTAL' in text_upper(t) and all(q not in text_upper(t) for q in ['QTY','QUANTITY','PCS']))

    subtotal_val = nearest_right_amount(subtotal_anchor) if subtotal_anchor else None
    total_val = nearest_right_amount(total_anchor) if total_anchor else None
    # VAT: avoid picking percentages like '5%'; search right tokens without %
    def nearest_right_vat(anchor, max_dx=0.6, y_tol=0.08):
        if not anchor:
            return None
        ax1 = anchor['x1']
        ay = anchor['yc']
        candidates = []
        for w in words:
            if w['x0'] >= ax1 and abs(w['yc'] - ay) <= y_tol and (w['x0'] - ax1) <= max_dx and in_amount_column(w):
                if '%' in w['text']:
                    continue
                if is_amount_token(w['text']):
                    val = parse_amount_token(w['text'])
                    if val is not None:
                        candidates.append((abs(w['yc'] - ay), w['x0'], val))
        if not candidates:
            # Try a small vertical window below the anchor (same column region)
            for w in words:
                if w['x0'] >= ax1 and (0 < (w['yc'] - ay) <= 0.06) and (w['x0'] - ax1) <= max_dx and in_amount_column(w):
                    if '%' in w['text']:
                        continue
                    if is_amount_token(w['text']):
                        val = parse_amount_token(w['text'])
                        if val is not None:
                            candidates.append((abs(w['yc'] - ay), w['x0'], val))
            if not candidates:
                return None
        # Pick closest by vertical distance, then rightmost by x
        candidates.sort(key=lambda t: (t[0], -t[1]))
        return candidates[0][2]

    # Consider all VAT anchors; choose the one nearest to totals region and with a valid amount to the right
    def all_vat_anchors():
        bad_terms = ['TRN', 'REG', 'REGISTRATION', 'INCLUSIVE']
        for w in words:
            u = text_upper(w['text'])
            if 'VAT' in u and not any(bt in u for bt in bad_terms):
                yield w

    candidate_vats = []
    for va in all_vat_anchors():
        vv = nearest_right_vat(va)
        if vv is not None and vv > 0:
            # rank by closeness to total anchor if present, else by y position (prefer lower on page)
            if total_anchor:
                dist_to_total = abs(va['yc'] - total_anchor['yc'])
            else:
                dist_to_total = 1.0 - va['yc']
            # Filter implausible VATs: VAT should be a small fraction of total (e.g., <= 30%)
            if total_val is not None and vv > 0.3 * total_val:
                continue
            candidate_vats.append((dist_to_total, va['yc'], vv))
    vat_val = None
    if candidate_vats:
        candidate_vats.sort(key=lambda t: (t[0], -t[1]))  # nearest to total, then lowest on page
        vat_val = candidate_vats[0][2]
    # total_val already computed above

    def fmt(v):
        return f"{v:.2f}" if isinstance(v, (int,float)) else "Not Found"

    return fmt(subtotal_val), fmt(vat_val), fmt(total_val)

def extract_amount_before_tax(lines, vat_amount_str: str, total_amount_str: str):
    """Estimate subtotal using VAT and Total when possible; fallback to heuristics.
    Avoid confusing quantities with amounts by preferring lines containing currency hints.
    """
    total_val = parse_number(total_amount_str) if total_amount_str else None
    vat_val = parse_number(vat_amount_str) if vat_amount_str else None

    # If VAT and Total are present and reasonable, compute subtotal
    if total_val is not None and vat_val is not None and total_val >= vat_val:
        candidate = total_val - vat_val
        # sanity check: positive and not absurd
        if candidate >= 0:
            return f"{candidate:.2f}"

    # Heuristic fallback: find a line that looks like subtotal/sub-total
    for line in lines:
        u = line.upper()
        if any(k in u for k in ['SUBTOTAL', 'SUB-TOTAL', 'AMOUNT BEFORE TAX', 'BEFORE VAT']):
            nums = re.findall(r'AED\s*[\d,.]+|[\d,.]+', line)
            if nums:
                val = parse_number(nums[-1])
                if val is not None:
                    return f"{val:.2f}"

    # As a last resort, pick the largest monetary number that is not the total
    monetary = []
    for line in lines:
        if any(k in line.upper() for k in ['AED', 'AMOUNT', 'VALUE', 'TOTAL']):
            for m in re.findall(r'AED\s*[\d,.]+|[\d,.]+', line):
                v = parse_number(m)
                if v is not None:
                    monetary.append(v)
    if monetary:
        monetary.sort(reverse=True)
        for v in monetary:
            if total_val is None or abs(v - total_val) > 1e-2:
                return f"{v:.2f}"

    return "Not Found"

def extract_vat_amount(lines):
    """
    Extract VAT amount by looking for values exactly in front of or below "VAT 5%" patterns
    """
    def safe_float(num_str):
        try:
            return float(num_str.replace(',', ''))
        except (ValueError, AttributeError):
            return float('inf')

    vat_value = None
    last_vat_line_idx = -1
    # Try to estimate total to bound VAT candidates
    total_guess = None
    for line in lines:
        u = line.upper()
        if any(k in u for k in ['GRAND TOTAL', 'TOTAL AMOUNT', 'AMOUNT PAYABLE', 'INVOICE TOTAL']):
            m = re.search(r'(?:AED|DHS|DIRHAM)\s*([\d,.]+)', line)
            if m:
                total_guess = parse_number(m.group(1))
                break
            nums = re.findall(r'\b(\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)\b', line)
            if nums:
                total_guess = parse_number(nums[-1])
                break
    for i, line in enumerate(lines):
        u = line.upper()
        if 'VAT' in u and not any(x in u for x in ['TRN', 'REG', 'REGISTRATION', 'INCLUSIVE']):
            last_vat_line_idx = i
            # 1) AED-marked amount on same line
            aed_match = re.search(r'(?:AED|DHS|DIRHAM)\s*([\d,.]+)', line, flags=re.IGNORECASE)
            if aed_match:
                cand = parse_number(aed_match.group(1))
                if cand is not None and (total_guess is None or cand <= 0.3 * total_guess):
                    vat_value = aed_match.group(1)
                    continue
            # 2) Non-percent numeric on same line
            candidates = re.findall(r'(?<!\d)(\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)(?!\s*%)', line)
            if candidates:
                # choose last number but ensure plausible vs total
                for c in reversed(candidates):
                    cand = parse_number(c)
                    if cand is not None and (total_guess is None or cand <= 0.3 * total_guess):
                        vat_value = c
                        break
                if vat_value is not None:
                    continue
            # 3) Look down 1-3 lines for amount (skip percent)
            for j in range(i+1, min(i+4, len(lines))):
                next_line = lines[j].strip()
                aed_match = re.search(r'(?:AED|DHS|DIRHAM)\s*([\d,.]+)', next_line, flags=re.IGNORECASE)
                if aed_match:
                    cand = parse_number(aed_match.group(1))
                    if cand is not None and (total_guess is None or cand <= 0.3 * total_guess):
                        vat_value = aed_match.group(1)
                        break
                candidates = re.findall(r'(?<!\d)(\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)(?!\s*%)', next_line)
                if candidates:
                    for c in reversed(candidates):
                        cand = parse_number(c)
                        if cand is not None and (total_guess is None or cand <= 0.3 * total_guess):
                            vat_value = c
                            break
                    if vat_value is not None:
                        break
    if vat_value is not None:
        return vat_value
    
    return "Not Found"

def extract_total_amount(lines):
    """Extract the final total amount payable, avoiding intermediate amounts and quantities.
    
    Priority order:
    1) Lines with explicit final total keywords (GRAND TOTAL, NET PAYABLE, AMOUNT PAYABLE, etc.)
    2) Lines with 'TOTAL' and currency, but only if they appear to be final totals
    3) The largest monetary value that appears to be a final amount (not subtotal, not VAT-related)
    """
    # Keywords that indicate final total amounts
    final_total_keywords = [
        'GRAND TOTAL', 'NET PAYABLE', 'NET AMOUNT', 'AMOUNT PAYABLE', 'TOTAL AMOUNT',
        'AMOUNT DUE', 'BALANCE DUE', 'INVOICE TOTAL', 'FINAL TOTAL', 'TOTAL PAYABLE'
    ]
    
    # Keywords that indicate intermediate amounts (avoid these)
    intermediate_keywords = ['SUBTOTAL', 'SUB-TOTAL', 'BEFORE VAT', 'BEFORE TAX', 'EXCLUDING VAT']
    
    # Keywords that indicate quantities (avoid these)
    quantity_keywords = ['QTY', 'QUANTITY', 'PCS', 'PIECES', 'ITEMS', 'TOTAL QTY', 'TOTAL PCS']
    
    def is_intermediate_amount(line: str) -> bool:
        """Check if line contains intermediate amount indicators"""
        u = line.upper()
        return any(k in u for k in intermediate_keywords)
    
    def is_quantity_line(line: str) -> bool:
        """Check if line contains quantity indicators"""
        u = line.upper()
        return any(k in u for k in quantity_keywords)
    
    def extract_monetary_values(line: str) -> list[float]:
        """Extract monetary values from a line, preferring currency-marked values"""
        values = []
        
        # First priority: currency-marked values (AED, DHS, DIRHAM)
        currency_patterns = [
            r'(?:AED|DHS|DIRHAM)\s*([\d,]+\.?\d*)',
            r'([\d,]+\.?\d*)\s*(?:AED|DHS|DIRHAM)'
        ]
        
        for pattern in currency_patterns:
            matches = re.findall(pattern, line, re.IGNORECASE)
            for match in matches:
                val = parse_number(match)
                if val is not None and val > 0:
                    values.append(val)
        
        # Second priority: plain numbers that look like monetary amounts
        if not values:
            # Look for numbers with decimal points or commas (typical for money)
            number_matches = re.findall(r'\b(\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)\b', line)
            for match in number_matches:
                val = parse_number(match)
                if val is not None and val > 10:  # Filter out very small numbers
                    values.append(val)
        
        return values
    
    # Strategy 1: Look for explicit final total keywords
    for line in lines:
        u = line.upper()
        if any(k in u for k in final_total_keywords) and not is_quantity_line(line) and 'IN WORDS' not in u:
            values = extract_monetary_values(line)
            if values:
                return f"{max(values):.2f}"
    
    # Strategy 2: Look for 'TOTAL' lines that appear to be final totals
    # (avoid lines with 'SUBTOTAL', 'BEFORE VAT', etc.)
    for line in lines:
        u = line.upper()
        if ('TOTAL' in u and 
            not is_intermediate_amount(line) and 
            not is_quantity_line(line) and 
            'IN WORDS' not in u and
            ('AED' in u or 'DHS' in u or 'DIRHAM' in u)):
            
            values = extract_monetary_values(line)
            if values:
                return f"{max(values):.2f}"
    
    # Strategy 3: Find the largest monetary value that appears to be a final amount
    # Look for lines that mention amounts but don't have intermediate indicators
    final_amount_candidates = []
    
    for line in lines:
        u = line.upper()
        if (is_intermediate_amount(line) or 
            is_quantity_line(line) or 
            'IN WORDS' in u or
            'VAT' in u):
            continue
        
        # Only consider lines that seem to be about final amounts
        if any(k in u for k in ['TOTAL', 'AMOUNT', 'PAYABLE', 'DUE', 'FINAL']):
            values = extract_monetary_values(line)
            if values:
                final_amount_candidates.extend(values)
    
    if final_amount_candidates:
        # Return the largest value as it's most likely the final total
        return f"{max(final_amount_candidates):.2f}"
    
    return "Not Found"

def save_to_excel(results):
    """Save results to Excel file"""
    try:
        df = pd.DataFrame(list(results.items()), columns=['Field', 'Value'])
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"multi_invoice_table_{timestamp}.xlsx"
        filepath = os.path.join(OUTPUT_FOLDER, filename)
        df.to_excel(filepath, index=False)
        print(f"Excel file saved: {filepath}")
        return filename
    except Exception as e:
        print(f"Error saving Excel file: {e}")
        return None

def save_to_word(results):
    """Save results to Word file"""
    try:
        doc = Document()
        doc.add_heading('Invoice Data Extraction Results', 0)
        
        # Add timestamp
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Add results table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Value'
        
        # Data rows
        for field, value in results.items():
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"multi_invoice_table_{timestamp}.docx"
        filepath = os.path.join(OUTPUT_FOLDER, filename)
        doc.save(filepath)
        print(f"Word file saved: {filepath}")
        return filename
    except Exception as e:
        print(f"Error saving Word file: {e}")
        return None

def save_table_to_excel(rows: list[dict]):
    """Save a list of row dicts to Excel as a table (one invoice per row)."""
    try:
        if not rows:
            return None
        columns = ['Page', 'Company Name', 'Invoice Number', 'Date', 'Seller TRN', 'Buyer TRN', 'VAT Amount', 'Total Amount']
        # Ensure all keys exist
        normalized = [{col: row.get(col, '') for col in columns} for row in rows]
        df = pd.DataFrame(normalized, columns=columns)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"multi_invoice_table_{timestamp}.xlsx"
        filepath = os.path.join(OUTPUT_FOLDER, filename)
        df.to_excel(filepath, index=False)
        print(f"Excel file saved: {filepath}")
        return filename
    except Exception as e:
        print(f"Error saving table to Excel: {e}")
        return None

def save_table_to_word(rows: list[dict]):
    """Save a list of row dicts to Word as a table."""
    try:
        if not rows:
            return None
        columns = ['Page', 'Company Name', 'Invoice Number', 'Date', 'Seller TRN', 'Buyer TRN', 'VAT Amount', 'Total Amount']
        doc = Document()
        doc.add_heading('Multi-Invoice Extraction Results', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        for i, col in enumerate(columns):
            table.rows[0].cells[i].text = col
        for row in rows:
            tr = table.add_row().cells
            for i, col in enumerate(columns):
                tr[i].text = str(row.get(col, ''))
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"multi_invoice_table_{timestamp}.docx"
        filepath = os.path.join(OUTPUT_FOLDER, filename)
        doc.save(filepath)
        print(f"Word file saved: {filepath}")
        return filename
    except Exception as e:
        print(f"Error saving table to Word: {e}")
        return None

if __name__ == "__main__":
    print("Starting invoice processing...")
    success = process_invoice()
    if success:
        print("✅ Invoice processing completed successfully!")
    else:
        print("❌ Invoice processing failed!")
